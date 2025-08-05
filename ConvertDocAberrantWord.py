# Copyright (c) 2025 TrustTJM
# 根据 Apache License 2.0 许可证授权
# 详情请参阅 LICENSE 文件

import os
import json
import shutil
import platform
from pathlib import Path
from docx import Document
from win32com import client as wc
from fontTools.ttLib import TTFont
from dataclasses import dataclass

# 是否输出完整日志
__mCurrentLogBufferIndex = -1
__mLogBufferCount = 1024
__mLogBuffer = [None] * __mLogBufferCount
__mLogFile = None

def InitLog() :
    global __mCurrentLogBufferIndex, __mLogFile
    __mCurrentLogBufferIndex = -1
    __mLogFile = open("日志.txt", "w", -1, "utf-8")

def WriteLog(aLogContent, aIsPrint = True) :
    global __mCurrentLogBufferIndex

    if aIsPrint :
        print(aLogContent)

    if __mCurrentLogBufferIndex < __mLogBufferCount :
        # 先写入缓冲区
        __mLogBuffer[__mCurrentLogBufferIndex] = aLogContent
        __mCurrentLogBufferIndex = __mCurrentLogBufferIndex + 1
    else :
        # 缓冲区满了, 写入文件并重置缓冲区索引
        __mLogFile.write("".join(__mLogBuffer))
        __mCurrentLogBufferIndex = -1

def FlushLog() :
    global __mCurrentLogBufferIndex, __mLogFile

    if __mCurrentLogBufferIndex >= 0 :
        __mLogFile.write("".join(__mLogBuffer[:__mCurrentLogBufferIndex]))
        __mCurrentLogBufferIndex = -1
    __mLogFile.close()
    __mLogFile = None


# dict[str(异常字), WordMappingInfo(字的映射信息)]
__mWordMappingInfoForAberrantWordDict = {}
# dict[str(字形), WordMappingInfo(字的映射信息)]
__mWordMappingInfoForGlyphNameDict = {}
def InitAberrantWordMapping(aIsUseDefaultMapping = True, aCustomAberrantWordMappingDict = None, aCustomFontDirList = None) :
    """初始化异常字对应的正常字的映射关系

    Args:
        aIsUseDefaultMapping (bool): 是否使用默认的映射关系, 即:已经生成好的
        aCustomAberrantWordMappingDict (dict[str, str]): 自定义异常字的映射关系字典, 如:{"一" : "一", "二" : "二"}
        aCustomFontDirList (list[str]): 自定义字体目录的列表
    """

    WriteLog(f"----------------开始处理异常字与正常字的映射关系\n")

    __mWordMappingInfoForAberrantWordDict.clear()
    __mWordMappingInfoForGlyphNameDict.clear()

    vOutputMappingJsonPath = "异常字对应正常文字映射关系.txt"

    vIsDefaultMappingValid = False
    if aIsUseDefaultMapping :
        WriteLog(f"尝试加载默认映射关系:{vOutputMappingJsonPath}\n")
        try:
            # 尝试从文件加载映射关系
            with open(vOutputMappingJsonPath, 'r', encoding='utf-8') as f:
                vDefaultMappingDict = json.load(f)
                # WriteLog(f"vDefaultMappingDict:{vDefaultMappingDict}\n")
                for vAberrantWord, vDefaultMapping in vDefaultMappingDict.items() :
                    vWordMappingInfo = GetWordMappingInfo()
                    vWordMappingInfo.mGlyphName = vDefaultMapping["字形名"]
                    vWordMappingInfo.mAberrantUnicode = vDefaultMapping["异常字的Unicode"]
                    vWordMappingInfo.mAberrantWord = vDefaultMapping["异常字"]
                    vWordMappingInfo.mNormalUnicode = vDefaultMapping["正常字的Unicode"]
                    vWordMappingInfo.mNormalWord = vDefaultMapping["正常字"]
                    __mWordMappingInfoForAberrantWordDict[vAberrantWord] = vWordMappingInfo

                # 映射关系起码要大于一定数量才算正确
                vLength = len(__mWordMappingInfoForAberrantWordDict)
                if vLength > 10 :
                    vIsDefaultMappingValid = True
                    WriteLog(f"默认映射关系加载成功, 有{vLength}对\n")
        except Exception as e:
            __mWordMappingInfoForAberrantWordDict.clear()
            WriteLog(f"加载默认映射关系失败:{vOutputMappingJsonPath}, 原因::{str(e)}\n")
    
    # 加载默认映射关系失败, 重新生成映射关系
    if not vIsDefaultMappingValid :
        WriteLog(f"加载默认映射关系失败, 重新生成映射关系\n")
        # 处理系统自带字体的映射关系
        ProcessFontMapping(GetSystemFontDir())
        # 添加自定义字体的映射关系
        if aCustomFontDirList :
            for vCustomFontDir in aCustomFontDirList :
                ProcessFontMapping(vCustomFontDir)
        # 添加自定义映射表
        if aCustomAberrantWordMappingDict :
            for vAberrantWord, vNormalWord in aCustomAberrantWordMappingDict.items() :
                vWordMappingInfo = GetWordMappingInfo()
                vWordMappingInfo.mGlyphName = "自定义"
                vWordMappingInfo.mAberrantUnicode = -9999
                vWordMappingInfo.mAberrantWord = vAberrantWord
                vWordMappingInfo.mNormalUnicode = -9999
                vWordMappingInfo.mNormalWord = vNormalWord
                __mWordMappingInfoForAberrantWordDict[vAberrantWord] = vWordMappingInfo

        if len(__mWordMappingInfoForAberrantWordDict) <= 0 :
            WriteLog(f"----------------处理映射关系出现异常\n")
        else :
            vWordMappingInfoForAberrantWordDict = {}
            for vAberrantWord, vWordMappingInfo in __mWordMappingInfoForAberrantWordDict.items() :
                vWordMappingInfoDict = {}
                vWordMappingInfoForAberrantWordDict[vAberrantWord] = vWordMappingInfoDict
                vWordMappingInfoDict["字形名"] = vWordMappingInfo.mGlyphName
                vWordMappingInfoDict["异常字的Unicode"] = vWordMappingInfo.mAberrantUnicode
                vWordMappingInfoDict["异常字"] = vWordMappingInfo.mAberrantWord
                vWordMappingInfoDict["正常字的Unicode"] = vWordMappingInfo.mNormalUnicode
                vWordMappingInfoDict["正常字"] = vWordMappingInfo.mNormalWord
            with open(vOutputMappingJsonPath, "w", encoding="utf-8") as f :
                json.dump(vWordMappingInfoForAberrantWordDict, f, ensure_ascii=False, indent=2)
                WriteLog(f"----------------映射关系已保存至:{os.path.abspath(vOutputMappingJsonPath)}\n")


def GetSystemFontDir() -> str :
    """获取系统字体的目录

    Returns:
        str: 目录路径
    """
    system = platform.system()
    if system == "Windows" :
        return os.environ.get("WINDIR", "C:/Windows") + "/Fonts"
    elif system == "Darwin" :
        return "/Library/Fonts"
    elif system == "Linux" :
        user_fonts = os.path.expanduser("~/.local/share/fonts")
        if os.path.exists(user_fonts) :
            return user_fonts
        else :
            return "/usr/share/fonts"
    else :
        raise OSError(f"Unsupported OS: {system}")


# 异常字的Unicode和正常字的Unicode的大小关系
__mIsAberrantLessNormalForFont = {
    "AdobeHeitiStd-Regular".lower() : True,
    "AdobeSongStd-Light".lower() : True,
    "Deng".lower() : True,
    "Dengb".lower() : True,
    "Dengl".lower() : True,
    "malgun".lower() : False,
    "malgunbd".lower() : False,
}
# 待检查的异常字列表, 为了提升检测效率
__mReuseNeedCheckvGlyphNameList = []
def ProcessFontMapping(aFontDir) :
    """处理字体的映射关系, 如果一个字体有相同字形的2个以上Unicode的字体, 则进行处理
    """

    if not os.path.exists(aFontDir) :
        raise FileNotFoundError(f"字体目录不存在: {aFontDir}")

    vFontPaths = GetFilePaths(aFontDir, [".ttf", ".ttc", ".otf"])
    WriteLog(f"开始生成异常字映射关系, ({aFontDir})目录下, 共检测到{len(vFontPaths)}个目标字体文件\n")

    # 遍历所有字体, 提取有效的映射关系
    for vFontPath in vFontPaths :
        try :
            # 获取异常字的Unicode和正常字的Unicode的大小关系
            vFileName, vFilePathExtension = os.path.splitext(os.path.basename(vFontPath).lower())
            vIsAberrantLessNormal = __mIsAberrantLessNormalForFont.get(vFileName)
            if vIsAberrantLessNormal == None :
                continue
            WriteLog(f"(字体名:{vFileName})的映射结果{str(vIsAberrantLessNormal)}\n", False)

            vTTFont = TTFont(vFontPath)
            if vTTFont == None :
                continue
            vCmap = vTTFont.getBestCmap()
            if vCmap == None :
                continue
            
            __mReuseNeedCheckvGlyphNameList.clear()

            # 这里的原理是只处理字体中同时包含异常字和其对应的正常字
            for vUnicode, vGlyphName in vCmap.items() :
                vWord = chr(vUnicode)

                # 先检测异常字的字典
                vWordMappingInfo = __mWordMappingInfoForAberrantWordDict.get(vWord)
                if vWordMappingInfo :
                    # 已经处理的跳过
                        continue
                
                # 在检测字形的字典
                vWordMappingInfo = __mWordMappingInfoForGlyphNameDict.get(vGlyphName)
                if not vWordMappingInfo :
                    vWordMappingInfo = GetWordMappingInfo()
                    __mWordMappingInfoForGlyphNameDict[vGlyphName] = vWordMappingInfo

                # 默认先赋值异常
                if vWordMappingInfo.mAberrantWord == "" :
                    vWordMappingInfo.mAberrantUnicode = vUnicode
                    vWordMappingInfo.mAberrantWord = vWord
                    __mReuseNeedCheckvGlyphNameList.append(vGlyphName)
                else :
                    # 根据具体情况判断是否需要交换异常字和正常字的数据
                    if ((vIsAberrantLessNormal and vWordMappingInfo.mAberrantUnicode > vUnicode) or
                        (not vIsAberrantLessNormal and vWordMappingInfo.mAberrantUnicode < vUnicode)) :
                        vWordMappingInfo.mNormalUnicode = mAberrantUnicode
                        vWordMappingInfo.mAberrantUnicode = vUnicode
                        vWordMappingInfo.mNormalWord = mAberrantWord
                        vWordMappingInfo.mAberrantWord = vWord
                    else :
                        vWordMappingInfo.mNormalUnicode = vUnicode
                        vWordMappingInfo.mNormalWord = vWord

            # 检测异常字列表
            for vGlyphName in __mReuseNeedCheckvGlyphNameList :
                vWordMappingInfo = __mWordMappingInfoForGlyphNameDict.get(vGlyphName)
                if vWordMappingInfo.mAberrantWord != "" and vWordMappingInfo.mNormalWord != "" :
                    # 有效, 添加到异常字的字典
                    __mWordMappingInfoForAberrantWordDict[vWordMappingInfo.mAberrantWord] = vWordMappingInfo
                else :
                    # 无效, 移除字形的字典并回收
                    del __mWordMappingInfoForGlyphNameDict[vGlyphName]
                    RecycleWordMappingInfo(vWordMappingInfo)

        except Exception as e :
            WriteLog(f"这个异常可以不看, 因为可能有的字体不标准, 只看最终文档是否转换成功即可, 检测字体:{vFontPath}时发生异常:{str(e)}\n")
            continue


# 字的映射信息
@dataclass
class WordMappingInfo:
    mGlyphName: str = "" # 字形名
    mAberrantUnicode: int = -1 # 异常字的Unicode
    mAberrantWord: str = "" # 异常字
    mNormalUnicode: int = -1 # 正常字的Unicode
    mNormalWord: str = "" # 正常字

__mWordMappingInfoPool = []
def GetWordMappingInfo() -> list[int] :
    if len(__mWordMappingInfoPool) > 0 :
        return __mWordMappingInfoPool.pop()
    else :
        return WordMappingInfo()

def RecycleWordMappingInfo(aWordMappingInfo) :
    aWordMappingInfo.mGlyphName = ""
    aWordMappingInfo.mAberrantUnicode = -1
    aWordMappingInfo.mAberrantWord = ""
    aWordMappingInfo.mNormalUnicode = -1
    aWordMappingInfo.mNormalWord = ""
    __mWordMappingInfoPool.append(aWordMappingInfo)


def GetFilePaths(aDir, aExtensions) -> list[str] :
    """递归获取指定目录下的制定后缀名文件路径

    Args:
        aDir (str): 指定目录
        aExtensions (list[str]): 后缀名数组, 如:[".vDoc", ".docx"]

    Returns:
        list[str]: 文件路径列表
    """

    aExtensions = [vExtension.lower() for vExtension in aExtensions]
    vFilterFilePaths = []

    # WriteLog(f"GetFilePaths---1---aDir:{aDir}---aExtensions:{aExtensions}\n")
    for vRoot, vDirs, vFilePaths in os.walk(aDir) :
        for vFilePath in vFilePaths :
            vFileName, vFilePathExtension = os.path.splitext(vFilePath)
            # WriteLog(f"GetFilePaths---2---vFilePath:{vFilePath}, vFileName:{vFileName}, vFilePathExtension:{vFilePathExtension}, vFilePathExtension.lower():{vFilePathExtension.lower()}\n")
            if vFilePathExtension.lower() in aExtensions :
                # WriteLog(f"GetFilePaths---3---{os.path.join(vRoot, vFilePath)}\n")
                vFilterFilePaths.append(os.path.join(vRoot, vFilePath))

    return vFilterFilePaths


def TestSpecifiedAberrantWord() :
    """测视指定的异常字的调试函数
    """

    vFontDir = GetSystemFontDir()
    if not os.path.exists(vFontDir) :
        raise FileNotFoundError(f"系统字体目录不存在: {vFontDir}")

    vFontPaths = GetFilePaths(vFontDir, [".ttf", ".ttc", ".otf"])
    WriteLog(f"({vFontDir})目录下, 共检测到{len(vFontPaths)}个目标字体文件\n")

    # 遍历所有字体, 提取有效的映射关系
    for vFontPath in vFontPaths :
        try :
            vTTFont = TTFont(vFontPath)
            vCmap = vTTFont.getBestCmap()
            
            for vUnicode, vGlyphName in vCmap.items() :
                try :
                    vWord = chr(vUnicode)
                    
                    if (vWord == "⾕" or
                        vWord == "⽆" or
                        vWord == "了" or
                        vWord == "不" or
                        vWord == "⻘" or
                        vWord == "⼗" or
                        vWord == "里" or
                        vWord == "易" or
                        vWord == "—") :
                        WriteLog(f"异常文字---({vFontPath})---({vGlyphName})---({vUnicode})---({vWord})\n", False)
                    elif (vWord == "谷" or
                          vWord == "无" or
                          vWord == "了" or
                          vWord == "不" or
                          vWord == "青" or
                          vWord == "十" or
                          vWord == "里" or
                          vWord == "易" or
                          vWord == "一") :
                        WriteLog(f"正常文字---({vFontPath})---({vGlyphName})---({vUnicode})---({vWord})\n", False)
                except ValueError :
                    WriteLog(f"异常---({vFontPath})---({vGlyphName})---({vUnicode})---({str(vWord)})\n", False)
                    continue
        except Exception as e :
            WriteLog(f"检测字体:{vFontPath}时发生异常:{str(e)}\n")
            continue


def ConvertWordFlies(aWordDir = None) :
    """转换Word文档

    Args:
        aWordDir (str): Word文档所在的目录路径
    """

    vOutputDir = None
    if aWordDir == None :
        aWordDir = os.getcwd()
        vOutputDir = f"{aWordDir}/Output"
    else :
        vOutputDir = os.path.join(os.path.dirname(aWordDir), f"{os.path.basename(aWordDir)}_Output")

    # 获取需要转换的文件, 并创建导出的文件夹, 并把所有的文件复制过去, 避免直接修改源文件
    
    if os.path.exists(vOutputDir):
        shutil.rmtree(vOutputDir)
    os.makedirs(vOutputDir)
    
    vFilePaths = GetFilePaths(aWordDir, [".doc", ".docx"])
    for vFilePath in vFilePaths :
        vSrcPath = Path(vFilePath).resolve()
        vDstPath = Path(vOutputDir) / vSrcPath.relative_to(aWordDir)
        # 创建子目录
        if not os.path.exists(vDstPath.parent):
            vDstPath.parent.mkdir(True, True)
        shutil.copy2(vSrcPath, vDstPath)

        try:
            vDocxPath = vDstPath
            if vDstPath.suffix.lower() == ".doc" :
                # 将.doc文件转换为.docx格式
                vDocxPath = vDstPath.replace(".doc", ".docx")
                vDocFile = wc.Dispatch("Word.Application")
                vDoc = vDocFile.Documents.Open(doc_path)
                vDoc.SaveAs(vDocxPath, 16)  # 16代表.docx格式
                vDoc.Close()
                vDocFile.Quit()
                os.remove(vDstPath)
                WriteLog(f"发现.doc文件, 将其转换为.docx; {vDstPath} -> {vDocxPath}\n")

            # 替换文本
            ReplaceDocText(vDocxPath)
        except Exception as e:
            WriteLog(f"转换 {vDstPath} 时出错:{str(e)}\n")
    
__mReuseAberrantWordLogList = []
__mReuseNormalWordLogList = []
def ReplaceRunText(aKeyLog, aRun) :
    vRunText = None
    if len(__mReuseAberrantWordLogList) > 0 :
        __mReuseAberrantWordLogList.clear()
        __mReuseNormalWordLogList.clear()
    for vAberrantWord, vWordMappingInfo in __mWordMappingInfoForAberrantWordDict.items() :
        if vAberrantWord in aRun.text :
            if vRunText == None :
                vRunText = aRun.text
            aRun.text = aRun.text.replace(vAberrantWord, vWordMappingInfo.mNormalWord)
            __mReuseAberrantWordLogList.append(vAberrantWord)
            __mReuseNormalWordLogList.append(vWordMappingInfo.mNormalWord)

    if len(__mReuseAberrantWordLogList) > 0 :
        WriteLog(f"在({aKeyLog})找到异常字\n---异常字:({", ".join(__mReuseAberrantWordLogList)})\n---正常字:({", ".join(__mReuseNormalWordLogList)})\n---{aKeyLog}:({vRunText})\n---替换后:({aRun.text})\n", False)

def ReplaceDocText(aFilePath) :
    """替换Word文档中文本

    Args:
        aFilePath (str): Word文档路径
    """

    WriteLog(f"----------------开始转换:{aFilePath}\n")

    vDoc = Document(aFilePath)
    vLastMatchText = None
    
    # 转换段落
    for vParagraph in vDoc.paragraphs :
        for vRun in vParagraph.runs :
            ReplaceRunText("段落原文", vRun)
    
    # 转换表格
    for vTable in vDoc.tables :
        for vRow in vTable.rows :
            for vCell in vRow.cells :
                for vParagraph in vCell.paragraphs :
                    for vRun in vParagraph.runs :
                        ReplaceRunText("表格", vRun)
    
    # 转换页眉页脚
    for vSection in vDoc.sections :
        for vParagraph in vSection.header.paragraphs :
            for vRun in vParagraph.runs :
                ReplaceRunText("页眉", vRun)
        for vParagraph in vSection.footer.paragraphs :
            for vRun in vParagraph.runs :
                ReplaceRunText("页脚", vRun)
    
    vDoc.save(aFilePath)
    WriteLog(f"----------------转换完成:{aFilePath}\n")


# ----------------正式执行----------------
try:
    WriteLog(f"--------------------------------开始处理--------------------------------\n")
    InitLog()

    # 调试函数
    # TestSpecifiedAberrantWord()

    # 示例1:初始化异常字的映射关系; 参数1:不使用默认的映射关系; 参数2:自定义映射关系; 参数3:自定义字体的文件夹路径列表
    # InitAberrantWordMapping(False
    #                         {"一" : "一", "二" : "二"},
    #                         {"d:/自定义字体文件夹1", "d:/自定义字体文件夹2"})

    # 实例1:初始化异常字的映射关系; 使用默认我已经生成好的映射关系, 如果中间有问题, 则添加一个自定义关系再进行生成
    InitAberrantWordMapping(True, {"—" : "一"})

    # 实例2:初始化异常字的映射关系; 参数1:不使用默认我已经生成好的映射关系; 参数2:添加了一个映射关系
    # InitAberrantWordMapping(False, {"—" : "一"})

    # 示例:转换Word文档, 参数1:Word文档所在的目录路径, 什么都不填:表示要转换的文件和自定义字体都在与当前.py代码在同一个文件夹, 不需要自定义映射关系
    ConvertWordFlies()

    # 实例:转换Word文档, 转换后的新Word文档在源目录的Output文件夹下
    # ConvertWordFlies("D:/文档")

except Exception as e:
    WriteLog(f"出现异常:{str(e)}\n")
    
finally :
    WriteLog(f"--------------------------------处理完毕--------------------------------\n")
    FlushLog()